from docx import Document

def replace_text_in_paragraph(paragraph):
    inline = paragraph.runs
    for item in inline:
        txt_list = item.text.split()
        new_list = []
        for word in txt_list:
            if word != "" and word !=" " and word != ".":
                word += "@"   
            if ",@" in word:
                word = word.replace(",@", "@,")
            if ".@" in word:
                word = word.replace(".@", "@.")         
            new_list.append(word)
        tt= " ".join(new_list)
        
        item.text = item.text.replace(item.text, tt)
        print(item.text)
                   
        

template_file_path = 'demo.docx'
output_file_path = 'result.docx'
# output_file_path_middle = 'middle.docx'


template_document = Document(template_file_path)
# template_document.save(output_file_path_middle)
for paragraph in template_document.paragraphs:
    replace_text_in_paragraph(paragraph)

for table in template_document.tables:
        for col in table.columns:
            for cell in col.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)

template_document.save(output_file_path)


